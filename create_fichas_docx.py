# Script para criar três ficheiros .docx com as fichas de auto-compaixão
from docx import Document
from docx.shared import Pt


def add_title(doc, text):
    heading = doc.add_heading(text, level=1)
    return heading


def add_section(doc, title, content):
    if title:
        doc.add_heading(title, level=2)
    for line in content.split("\n"):
        doc.add_paragraph(line)


def save_doc(title, filename, blocks):
    doc = Document()
    # Ajuste simples de estilo: fonte base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    add_title(doc, title)
    for heading, text in blocks:
        add_section(doc, heading, text)
    doc.save(filename)
    print('Gravado:', filename)


ficha1_blocks = [
    ("Sessão / Data / Paciente / Terapeuta", "Sessão: _______   Data: _______   Paciente: ____________________   Terapeuta: ____________________"),
    ("Abordagem / Objetivo / Duração", "Abordagem: CBT integrada com autocompaixão\nObjetivo: Identificar pensamentos automáticos autocríticos e gerar uma resposta cognitiva mais equilibrada e compassiva.\nDuração estimada: 20–30 min (sessão); 10–15 min (tarefa)."),
    ("1) Situação", "Descreve a situação concreta (quando, onde, quem estava presente):\n................................................................................................................................................................................"),
    ("2) Pensamento automático", "Pensamento automático (escreve a frase tal como surgiu):\n................................................................................................................................................................................"),
    ("3) Emoções", "- Emoção principal: ______________ Intensidade: __ /100\n- Outras emoções: __________________________________________________"),
    ("4) Evidências a favor", "- 1) ........................................................................................................\n- 2) ........................................................................................................"),
    ("5) Evidências contra", "- 1) ........................................................................................................\n- 2) ........................................................................................................"),
    ("6) Voz compassiva", "Voz compassiva — imagina uma pessoa empática; anota o que diria (frase curta):\n................................................................................................................................................................................"),
    ("7) Reescrita compassiva", "Reescrita compassiva (frase alternativa verdadeira e acolhedora):\n................................................................................................................................................................................"),
    ("8) Observação corporal", "Observação corporal breve (antes/depois): onde sentiste tensão? mudou algo após a reescrita?\n................................................................................................................................................................................"),
    ("Tarefa", "Tarefa entre sessões (7 dias): regista 2 episódios por dia usando esta ficha. Levar à próxima sessão."),
    ("Notas para o terapeuta", "Oferecer modelos de frases (ex.: “Isto foi difícil — reconheço o teu esforço”; “Posso aprender com isto sem me destruir”). Em teleconsulta, o paciente pode preencher em Google Docs partilhado; terapeuta pode gravar a frase compassiva em áudio.")
]

ficha2_blocks = [
    ("Sessão / Data / Paciente / Terapeuta", "Sessão: _______   Data: _______   Paciente: ____________________   Terapeuta: ____________________"),
    ("Abordagem / Objetivo / Duração", "Abordagem: Compassion‑Focused Therapy (CFT)\nObjetivo: Treinar o sistema calmante via respiração e imagética compassiva.\nDuração estimada: 25–40 min; prática diária 5–15 min."),
    ("Breve psicoeducação", "“Temos três sistemas: ameaça, impulso e calmante. Vamos praticar o calmante para reduzir autocrítica.”"),
    ("A) Avaliação rápida", "Sinais de ativação (pensamentos/sensações): ................................................................\nSituações que mais ativam a autocrítica: ................................................................"),
    ("B) Respiração ritmada", "Inspira 4 seg — expira 6–8 seg. Anotar sensação inicial e final:\nAntes: ____________________   Depois: ____________________"),
    ("C) Imagética compassiva", "Figura compassiva: quem é? (real / imaginada / mistura): ______________________\nQualidades (3): 1) ______ 2) ______ 3) ______\nFrase recebida da figura (curta): ____________________"),
    ("D) Integração", "Criar uma “Auto‑Compassiva” (nome/âncora): ________________\nFrase âncora para usar em crise: ____________________"),
    ("Prática diária", "- [ ] Respiração ritmada (5 min)   - [ ] Imagética guiada (5–10 min)   - [ ] Repetir frase âncora quando crítico"),
    ("Notas para o terapeuta", "Evitar imagética intensa se houver história de trauma sem estabilidade; optar por figura distante/segura. Para teleconsulta: fornecer gravação áudio com a imagética.")
]

ficha3_blocks = [
    ("Sessão / Data / Paciente / Terapeuta", "Sessão: _______   Data: _______   Paciente: ____________________   Terapeuta: ____________________"),
    ("Abordagem / Objetivo / Duração", "Abordagem: Mindful Self‑Compassion\nObjetivo: Observar a autocrítica com presença e responder com cuidado corporal e palavras compassivas.\nDuração estimada: 10–20 min; prática diária 5–15 min."),
    ("1) Preparação", "Posição: sentado com pés no chão. Mãos: __________________. Nota corporal inicial: ____________________"),
    ("2) Respiração", "Atenção à respiração (2–3 respirações atentas): notar sem alterar."),
    ("3) Episódio de autocrítica", "Trazer à mente um episódio recente de autocrítica (descrição curta):\n............................................................................................................................................"),
    ("4) Localizar sensação", "Localizar a sensação no corpo: ...................................................................."),
    ("5) Toque de acolhimento", "Toque de acolhimento (opcional): colocar a mão no peito/estômago — descreve sensação:\n............................................................................................................................................"),
    ("6) Frases compassivas", "Frases exemplo:\n- “Isto é difícil agora.”\n- “Posso ser gentil comigo mesmo(a).”\n- “Estou a aprender como lidar com isto.”\nFrases escolhidas: 1) ____ 2) ____ 3) ____"),
    ("7) Encerramento", "O que mudou? (emocional/corporal):\n............................................................................................................................................"),
    ("Tarefa", "Tarefa diária: sempre que surgir autocrítica, parar 1 respiração + frase compassiva; registar 1 episódio por dia."),
    ("Notas para o terapeuta", "Se o toque ativa trauma, sugerir alternativa (mãos nas pernas, segurar objeto). Em teleconsulta, enviar áudio curto (3–6 min) com a prática guiada.")
]

if __name__ == "__main__":
    base_path = r"C:\Users\diogo\AAA\Projects\Exercises and Evaluation"
    save_doc("Ficha 1 — Reescrita Compassiva de Pensamentos", base_path + "\\Ficha_Reescrita_Compassiva.docx", ficha1_blocks)
    save_doc("Ficha 2 — Imagética Compassiva (CFT)", base_path + "\\Ficha_Imagem_Compassiva_CFT.docx", ficha2_blocks)
    save_doc("Ficha 3 — Prática Mindful Self‑Compassion", base_path + "\\Ficha_Practica_Mindful_SelfCompassion.docx", ficha3_blocks)
